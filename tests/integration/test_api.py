"""httpx.AsyncClient against the real FastAPI app (app/api/app.py), real
Postgres/Qdrant/Ollama underneath. `app/api/dependencies.py`'s `lru_cache`
singletons and `app/db.get_session`'s process-lifetime engine are
overridden with this test session's own per-test fixtures for the same
reason `tests/conftest.py`'s `db_session` doesn't reuse `app.db.base`
directly: they'd otherwise bind async connections to one test's event
loop and break on the next test's fresh one.
"""

from __future__ import annotations

import uuid
from collections.abc import AsyncIterator

import pytest
import pytest_asyncio
from httpx import ASGITransport, AsyncClient

from app.api.app import app
from app.api.dependencies import (
    get_dense_client,
    get_generation_client,
    get_generic_chunker,
    get_reranker,
    get_sparse_embedder,
    get_standard_chunker,
    get_standard_store,
)
from app.db import get_session

pytestmark = pytest.mark.integration


@pytest_asyncio.fixture
async def api_client(
    db_session,
    dense_client,
    sparse_embedder,
    generic_chunker,
    standard_chunker,
    standard_store,
    reranker,
    generation_client,
) -> AsyncIterator[AsyncClient]:
    async def _session_override():
        yield db_session

    app.dependency_overrides[get_session] = _session_override
    app.dependency_overrides[get_dense_client] = lambda: dense_client
    app.dependency_overrides[get_sparse_embedder] = lambda: sparse_embedder
    app.dependency_overrides[get_generic_chunker] = lambda: generic_chunker
    app.dependency_overrides[get_standard_chunker] = lambda: standard_chunker
    app.dependency_overrides[get_standard_store] = lambda: standard_store
    app.dependency_overrides[get_reranker] = lambda: reranker
    app.dependency_overrides[get_generation_client] = lambda: generation_client
    try:
        transport = ASGITransport(app=app)
        async with AsyncClient(transport=transport, base_url="http://test") as client:
            yield client
    finally:
        app.dependency_overrides.clear()


async def test_health(api_client: AsyncClient):
    response = await api_client.get("/health")
    assert response.status_code == 200
    assert response.json() == {"status": "ok"}


async def test_full_client_project_workflow(api_client: AsyncClient, unique_name: str, qdrant_client):
    response = await api_client.post("/clients", json={"name": unique_name})
    assert response.status_code == 200
    client_body = response.json()
    client_id = client_body["id"]
    collection_name = client_body["qdrant_collection_name"]
    assert collection_name.startswith("client_")

    try:
        response = await api_client.get("/clients")
        assert any(c["id"] == client_id for c in response.json())

        response = await api_client.get(f"/clients/{client_id}")
        assert response.status_code == 200
        assert response.json()["id"] == client_id

        response = await api_client.post(f"/clients/{client_id}/projects", json={"name": "Audit"})
        assert response.status_code == 200
        project_id = response.json()["id"]
        assert response.json()["status"] == "evidence_collection"

        response = await api_client.get(f"/clients/{client_id}/projects")
        assert len(response.json()) == 1

        response = await api_client.get(f"/projects/{project_id}")
        assert response.status_code == 200
        assert response.json()["id"] == project_id

        response = await api_client.put(
            f"/projects/{project_id}/requirements/8.4.2", json={"status": "in_progress", "notes": "draft"}
        )
        assert response.status_code == 200
        assert response.json()["status"] == "in_progress"

        # second PUT on the same requirement upserts, doesn't duplicate
        response = await api_client.put(
            f"/projects/{project_id}/requirements/8.4.2", json={"status": "compliant", "notes": None}
        )
        assert response.status_code == 200
        response = await api_client.get(f"/projects/{project_id}/requirements")
        assert len(response.json()) == 1
        assert response.json()[0]["status"] == "compliant"

        response = await api_client.post(
            f"/projects/{project_id}/findings",
            json={"requirement_id": "8.4.2", "description": "gap", "recommendation": "fix it"},
        )
        assert response.status_code == 200
        finding_id = response.json()["id"]
        assert response.json()["status"] == "draft"

        response = await api_client.get(f"/projects/{project_id}/findings")
        assert len(response.json()) == 1

        response = await api_client.get(f"/projects/{project_id}/findings/{finding_id}")
        assert response.status_code == 200
        assert response.json()["id"] == finding_id

        response = await api_client.patch(
            f"/projects/{project_id}/findings/{finding_id}", json={"status": "reviewed"}
        )
        assert response.status_code == 200
        assert response.json()["status"] == "reviewed"
        assert response.json()["description"] == "gap"  # untouched fields survive a partial PATCH

        # pagination: two more findings, limit=1 should page through all three
        for _ in range(2):
            await api_client.post(
                f"/projects/{project_id}/findings",
                json={"requirement_id": "8.4.2", "description": "another gap"},
            )
        seen_ids = set()
        for offset in range(3):
            page = await api_client.get(f"/projects/{project_id}/findings", params={"limit": 1, "offset": offset})
            assert len(page.json()) == 1
            seen_ids.add(page.json()[0]["id"])
        assert len(seen_ids) == 3

        response = await api_client.post(
            f"/projects/{project_id}/evidence",
            files={"file": ("interview.txt", b"MFA is enforced via TOTP.", "text/plain")},
            data={"evidence_type": "interview_transcript"},
        )
        assert response.status_code == 200
        evidence = response.json()
        assert evidence["filename"] == "interview.txt"
        assert evidence["requirement_id"] is None

        response = await api_client.get(f"/projects/{project_id}/evidence")
        assert len(response.json()) == 1

        response = await api_client.get(f"/projects/{project_id}/evidence/{evidence['id']}")
        assert response.status_code == 200
        assert response.json()["id"] == evidence["id"]

        response = await api_client.put(
            f"/projects/{project_id}/evidence/{evidence['id']}/requirement", json={"requirement_id": "8.4.2"}
        )
        assert response.status_code == 200
        assert response.json()["requirement_id"] == "8.4.2"

        # DELETE evidence must remove its Qdrant chunks, not just the row
        response = await api_client.delete(f"/projects/{project_id}/evidence/{evidence['id']}")
        assert response.status_code == 200
        points, _ = await qdrant_client.scroll(
            collection_name=collection_name,
            scroll_filter={
                "must": [{"key": "document_id", "match": {"value": evidence["qdrant_document_id"]}}]
            },
            limit=10,
        )
        assert points == []
        assert (await api_client.get(f"/projects/{project_id}/evidence/{evidence['id']}")).status_code == 404
    finally:
        # Dogfooding the new endpoint as this test's own cleanup: also
        # verifies DELETE /clients/{id} actually removes the collection.
        response = await api_client.delete(f"/clients/{client_id}")
        assert response.status_code == 200
        assert not await qdrant_client.collection_exists(collection_name)
        assert (await api_client.get(f"/clients/{client_id}")).status_code == 404


async def test_project_not_found_returns_404(api_client: AsyncClient):
    unknown = str(uuid.uuid4())
    assert (await api_client.get(f"/clients/{unknown}")).status_code == 404
    assert (await api_client.delete(f"/clients/{unknown}")).status_code == 404
    assert (await api_client.get(f"/projects/{unknown}")).status_code == 404
    assert (await api_client.delete(f"/projects/{unknown}")).status_code == 404
    assert (await api_client.get(f"/projects/{unknown}/requirements")).status_code == 404
    assert (await api_client.get(f"/projects/{unknown}/findings")).status_code == 404
    assert (await api_client.get(f"/projects/{unknown}/findings/{unknown}")).status_code == 404
    assert (await api_client.patch(f"/projects/{unknown}/findings/{unknown}", json={"status": "final"})).status_code == 404
    assert (await api_client.delete(f"/projects/{unknown}/findings/{unknown}")).status_code == 404
    assert (await api_client.get(f"/projects/{unknown}/evidence")).status_code == 404
    assert (await api_client.get(f"/projects/{unknown}/evidence/{unknown}")).status_code == 404
    assert (await api_client.delete(f"/projects/{unknown}/evidence/{unknown}")).status_code == 404
    assert (await api_client.post(f"/projects/{unknown}/chat", json={"question": "x"})).status_code == 404


async def test_project_chat_endpoint(api_client: AsyncClient, client_and_project):
    _client, project = client_and_project
    response = await api_client.post(f"/projects/{project.id}/chat", json={"question": "Что говорит стандарт про MFA?"})
    assert response.status_code == 200
    assert "answer" in response.json()


async def test_upload_knowledge_document(api_client: AsyncClient, standard_store, unique_name: str):
    import io

    import docx

    buffer = io.BytesIO()
    document = docx.Document()
    document.add_heading("Test Guidance Section", level=1)
    document.add_paragraph("This is a throwaway guidance document uploaded by a test.")
    document.save(buffer)
    buffer.seek(0)

    response = await api_client.post(
        "/documents",
        files={"file": (f"{unique_name}.docx", buffer.read(), "application/vnd.openxmlformats")},
        data={"document_type": "guidance"},
    )
    assert response.status_code == 200
    body = response.json()
    assert body["chunk_count"] > 0
    assert body["document_type"] == "guidance"

    try:
        points, _ = await standard_store.client.scroll(
            collection_name=standard_store.collection_name,
            scroll_filter={"must": [{"key": "document_id", "match": {"value": body["document_id"]}}]},
            limit=10,
        )
        assert len(points) == body["chunk_count"]
    finally:
        await standard_store.delete_by_document_id(body["document_id"])


async def test_upload_knowledge_document_unsupported_type_returns_400(api_client: AsyncClient):
    response = await api_client.post(
        "/documents",
        files={"file": ("notes.exe", b"not a real document", "application/octet-stream")},
        data={"document_type": "other"},
    )
    assert response.status_code == 400
