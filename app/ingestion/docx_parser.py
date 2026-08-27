"""Section extraction for organizational evidence documents
(`data/documents/org/`, e.g. an ISMS policy) - arbitrary .docx structure,
not the PCI DSS table template. Sections are read directly from Word's own
"Heading N" paragraph styles rather than a text heuristic, since a .docx
file already carries that structure natively. See ARCHITECTURE.md
section 14.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

import docx


@dataclass
class DocxSection:
    heading: str | None
    paragraphs: list[str] = field(default_factory=list)

    @property
    def text(self) -> str:
        return "\n".join(self.paragraphs)


class DocxParser:
    def __init__(self, path: Path) -> None:
        self.path = path

    def extract_sections(self) -> list[DocxSection]:
        document = docx.Document(self.path)
        sections: list[DocxSection] = []
        current = DocxSection(heading=None)

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = paragraph.style.name if paragraph.style else ""
            if style_name.startswith("Heading"):
                if current.heading or current.paragraphs:
                    sections.append(current)
                current = DocxSection(heading=text)
            else:
                current.paragraphs.append(text)

        if current.heading or current.paragraphs:
            sections.append(current)
        return sections

    def guess_title(self, sections: list[DocxSection]) -> str:
        for section in sections:
            if section.paragraphs:
                return section.paragraphs[0]
        return self.path.stem
